from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import pickle
import os
from datetime import datetime
from gspread_dataframe import get_as_dataframe
from openpyxl import Workbook
import gspread
from google.oauth2.service_account  import Credentials
# import google.oauth2.service_account # import Credentials

loaded_data =[]
_path = './project_doc' 
seen = []
try:
    # Load from a pickle file
    with open("data.pkl", "rb") as file:  # "rb" means read in binary mode
        loaded_data = pickle.load(file)

    for data in loaded_data:
        print(data)
        seen.append(data)

    # print("Loaded data:", loaded_data)
except FileNotFoundError:
    print("File does not exist")

sheet_configs = [
    {"spreadsheet_id": "1CNVUPciOW4OxKL5khCmL2sgVHNqpdBBkuKRjskrVaMs", "worksheet_name": "การตอบแบบฟอร์ม 1"}, # spreadsheet ID 1 ECE Capstone Project Evaluation Form
]

SCOPES = ["https://www.googleapis.com/auth/spreadsheets", 
          "https://www.googleapis.com/auth/drive"]
# # download credentials from JSON 
creds = Credentials.from_service_account_file("capstonemanagement-5cab3e556827.json", scopes=SCOPES)

# ใช้ gspread login
client = gspread.authorize(creds)
path = 'capstonemanagement-5cab3e556827.json'
  
# Check whether the specified path exists or not 
isExist = os.path.exists(path) 
# print(isExist) 

gc = gspread.service_account(filename = path)
sh = gc.open_by_key('1CNVUPciOW4OxKL5khCmL2sgVHNqpdBBkuKRjskrVaMs')
worksheet = sh.worksheet("การตอบแบบฟอร์ม 1")

data = worksheet.get_all_records()
df = pd.DataFrame(data)



# # df = pd.read_excel("student.xlsx", engine="openpyxl")
# df = pd.read_excel("Capstone design project.xlsx", engine="openpyxl")

for index, row in df.iterrows():
    # print(row['อาจารย์ที่ปรึกษาหลักปริญญานิพนธ์'])
    proj_code = row['รหัสโครงงาน']
    semester = row['ปีการศึกษา']
    subject = row['วิชา']
    chk = {"proj_code": proj_code, "semester": semester, "subject": row['วิชา']  }
    # print(chk)
    if chk not in seen:
        seen.append(chk)
        loaded_data.append(chk)

        # Create a new document
        doc = Document()
        # Title
        first_sem = [6, 7, 8, 9, 10]
        sec_sem = [11, 12, 1, 2, 3]
        summer_sem = [4, 5]
        ts_val = pd.to_datetime(row['ประทับเวลา'], dayfirst=True)
        month_val  = ts_val.month
        year_val = ts_val.year
        year_add = 543
        if month_val  in first_sem:
            semester_ = 'First Semester'
        elif month_val  in sec_sem:
            semester_ ="Second Semester"
            year_add = year_add -1 
        else: 
            semester_ ="Third Semester"
            year_add = year_add -1 

        # print("First Semester/", (year_val+year_add))
        # print(proj_code)


        title_txt = 'Capstone Design Project Exam: '+ str(year_val+year_add) #row['ปีการศึกษา']
        title = doc.add_heading(title_txt, level=1)
        run = title.runs[0]
        run.font.name = 'Times New Roman'  # Change font type
        run.font.size = Pt(24)  # Change font size
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Subtitle
        subtitleFontSize = 12 
        subtitle = doc.add_paragraph()
        subtitle.paragraph_format.space_after = Pt(3)
        subtitle.paragraph_format.before = Pt(0)
        subtitle = subtitle.add_run("Subject : "  + row['วิชา'])
        subtitle.font.name = 'Times New Roman'
        subtitle.font.size = Pt(subtitleFontSize)  # Change font size
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Subtitle
        Engsubtitle = doc.add_paragraph()
        Engsubtitle.paragraph_format.space_after = Pt(0)
        Engsubtitle.paragraph_format.space_before = Pt(0)
        Engsubtitle_txt = "Project Name: " + row['ชื่อโครงการ (ภาษาอังกฤษ)']
        Engsubtitle = Engsubtitle.add_run(Engsubtitle_txt)
        Engsubtitle.font.name = 'Times New Roman'
        Engsubtitle.font.size = Pt(subtitleFontSize)  # Change font size
        Engsubtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

#         THsubtitleFontSize = 14           
#         THsubtitle = doc.add_paragraph()
#         THsubtitle.paragraph_format.space_after = Pt(0)
#         THsubtitle.paragraph_format.space_before = Pt(0)
#         THsubtitle_txt = "ชื่อโปรเจ็ค: " + row['ชื่อโครงการ (ภาษาไทย)']
#         THsubtitle = THsubtitle.add_run(THsubtitle_txt)
#         THsubtitle.font.name = 'SarabunPSK'
#         THsubtitle.font.size = Pt(THsubtitleFontSize)  # Change font size
#         THsubtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Convert to datetime object
        dt = datetime.strptime(row['วันที่สอบ'], "%d/%m/%Y") #  "%Y-%m-%d %H:%M:%S"
        dt_day = dt.day
        dt_month = dt.month
        year_th = dt.year
        if dt.year==year_val :
            year_th = year_th + 543  # แปลง ค.ศ. เป็น พ.ศ.

        thai_date_str = f"{dt_day}/{dt_month}/{year_th}"
        Date_title = doc.add_paragraph()
        Date_title.paragraph_format.space_after = Pt(2)
        Date_title.paragraph_format.space_before = Pt(0)
        Date_title_txt = "Exam Date: " + thai_date_str
        Date_title = Date_title.add_run(Date_title_txt)
        Date_title.font.name = 'Times New Roman'
        Date_title.font.size = Pt(subtitleFontSize)  # Change font size
        Date_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("\n")  # Add spacing

        # Section: Member
        Member_heading = doc.add_heading("Member", level=2)
        run = Member_heading.runs[0]
        run.font.name = 'Times New Roman'  # Change font type
        run.font.size = Pt(14)  # Change font size
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Sample list with data
        data = [
            ["Section", "Student ID", "Student Name", "Signature"],
                [str(row['section ของคนที่ 1']), str(row['รหัสนักศึกษา คนที่ 1']), row['ชื่อ-นามสกุล สมาชิก   คนที่ 1'], ""],
                [str(row['section ของคนที่ 2']), str(row['รหัสนักศึกษา คนที่ 2']), row['ชื่อ-นามสกุล สมาชิก   คนที่ 2'], ""]
        ]

        if not pd.isna(row['section ของคนที่ 3']):
            data.append([str(row['section ของคนที่ 3']), str(row['รหัสนักศึกษา คนที่ 3']), row['ชื่อ-นามสกุล สมาชิก   คนที่ 3'], ""])

        # Define fonts for each column
        column_fonts = ["Times New Roman","Times New Roman", "SarabunPSK", "SarabunPSK"]
        column_sizes = [10, 10, 10, 10]

        # # Create a table with rows and columns based on data
        table = doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = "Table Grid"
        column_widths = [Inches(0.5), Inches(1.2), Inches(2.2), Inches(2.0)]
        for col_idx, width in enumerate(column_widths):
            for cell in table.columns[col_idx].cells:
                cell.width = width

        # # Populate table from list and apply fonts
        for row_idx, row_data in enumerate(data):  
            if row_idx == 0:
                column_fonts = ["Times New Roman","Times New Roman", "Times New Roman", "Times New Roman"]
            else:
                column_fonts = ["Times New Roman", "Times New Roman","SarabunPSK", "SarabunPSK"]

            for col_idx, value in enumerate(row_data):  
                cell = table.cell(row_idx, col_idx)  
                cell.text = value  # Set cell text
                
                # Apply font settings
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = column_fonts[col_idx]  # Set different font per column
                        run.font.size = Pt(column_sizes[col_idx])  # Set different font size
                        
        #                 # Ensure Word recognizes the font change
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), column_fonts[col_idx])


        # # Set row height and add space after each row
        for row_ in table.rows:
            # Set the row height
            row_.height = Pt(18)  # Adjust the height as needed

            # Add space after the row by adding an empty paragraph
            # This creates visual space between rows
            empty_paragraph = doc.add_paragraph()
            empty_paragraph.add_run()
            empty_paragraph_format = empty_paragraph.paragraph_format
            empty_paragraph_format.space_after = Pt(12)  # Adjust the space after as needed

        # message = "Submit Presentation : Yes" if row['submit report'] == "yes" else "Submit Presentation : No" 

        # add_Submit =  "Yes" if row['submit report'] == "Submit Presentation : Yes" else "Submit Presentation : No" 
        # doc.add_paragraph("\n")  # Add spacing
        # Submit_report = doc.add_paragraph()
        # Submit_report .paragraph_format.space_after = Pt(3)
        # Submit_report .paragraph_format.space_before = Pt(3)
        # Submit_report = Submit_report.add_run(add_Submit)
        # Submit_report.font.name = 'Times New Roman'
        # Submit_report.font.size = Pt(14)  # Change font size
        # Submit_report.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.add_paragraph("\n")  # Add spacing
        # Section: Committee 
        Committee_title = doc.add_paragraph()
        Committee_title .paragraph_format.space_after = Pt(0)
        Committee_title .paragraph_format.space_before = Pt(0)
        Committee_title = Committee_title.add_run("Committee Signature")
        Committee_title.font.name = 'Times New Roman'
        Committee_title.font.size = Pt(12)  # Change font size
        Committee_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        Committee_FontSize = 14
        Comittee_ls = [row['อาจารย์ที่ปรึกษาหลักปริญญานิพนธ์']] 

        if not pd.isna(row['อาจารย์ที่ปรึกษาร่วมปริญญานิพนธ์']) :
            Comittee_ls.append(row['อาจารย์ที่ปรึกษาร่วมปริญญานิพนธ์'])

        # Section: Committee 
        cnt = 0
        for i in range(len(Comittee_ls) + 4) :
            # print('i: ', i, 'cnt: ', cnt)
            if cnt == 0:
                add_commit = "ที่ปรึกษา : " + Comittee_ls[cnt] + "  ................................"
            else :
                if not pd.isna(row['อาจารย์ที่ปรึกษาร่วมปริญญานิพนธ์']) and cnt == 1:
                    add_commit = "ที่ปรึกษาร่วม : " + Comittee_ls[cnt] + "  ................................"
                else:
                    add_commit = "กรรมการ :                              ................................"

            Committee_title = doc.add_paragraph()
            Committee_title .paragraph_format.space_after = Pt(3)
            Committee_title .paragraph_format.space_before = Pt(3)
            Committee_title = Committee_title.add_run(add_commit)
            Committee_title.font.name = 'SarabunPSK'
            Committee_title.font.size = Pt(Committee_FontSize)  # Change font size
            Committee_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cnt+=1

        # Save the document
        file_name = str(proj_code) + ".docx"
        doc.save(os.path.join(_path, file_name))
        print("created: "+file_name)

with open("data.pkl", "wb") as file:  # "wb" means write in binary mode
    pickle.dump(loaded_data, file)
print("Data saved successfully!")